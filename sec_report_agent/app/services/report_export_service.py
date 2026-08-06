"""报告导出服务（V2.1）— 版本正文导出 Markdown / Word(docx)
md：直接返回 content_md；docx：解析 Markdown（标题/列表/表格/引用）生成 Word。
零新增依赖（python-docx 已随底座安装）。
"""

import re
from io import BytesIO

from common.logger.logger import LogManager

logger = LogManager.get_logger()


class ReportExportService:
    """报告导出：按 format 返回 bytes + 文件名"""

    @staticmethod
    def build_markdown(content_md: str) -> bytes:
        return content_md.encode("utf-8")

    @staticmethod
    def build_docx(content_md: str) -> bytes:
        """Markdown → docx（标题/引用/列表/表格/段落）"""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document()
        lines = content_md.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            if line.startswith("# "):
                h = doc.add_heading(line[2:], level=0)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                for run in p.runs:
                    run.font.size = Pt(9)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+[.、]", line):
                doc.add_paragraph(line, style="List Number")
            elif line.startswith("|"):
                # 收集连续表格行
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_rows.append(lines[i].strip())
                    i += 1
                doc = ReportExportService._fill_table(doc, table_rows)
                continue
            else:
                doc.add_paragraph(line)
            i += 1

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _fill_table(doc, rows: list[str]):
        """Markdown 表格行 → docx 表格（第二行分隔符跳过）"""
        from docx import Document  # noqa: F401  (type hint only)

        def cells(row: str) -> list[str]:
            body = row.strip().strip("|")
            return [c.strip() for c in body.split("|")]

        data = []
        for row in rows:
            cs = cells(row)
            if all(re.fullmatch(r":?-{3,}:?", c) for c in cs if c):
                continue  # 分隔符行
            data.append(cs)
        if not data:
            return doc
        ncols = max(len(r) for r in data)
        table = doc.add_table(rows=len(data), cols=ncols)
        table.style = "Light Grid Accent 1"
        for ri, row in enumerate(data):
            for ci in range(ncols):
                table.rows[ri].cells[ci].text = row[ci] if ci < len(row) else ""
        return doc
