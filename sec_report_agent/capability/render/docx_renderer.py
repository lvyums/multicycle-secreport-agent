"""Docx 渲染器 — python-docx 生成 Word 报告（标题/段落/表格）"""

import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from capability.render.render_base import Renderer
from model.struct.structs import RenderData
from common.logger.logger import LogManager

logger = LogManager.get_logger()


class DocxRenderer(Renderer):
    """Word 渲染器"""

    ext = "docx"

    def render(self, data: RenderData) -> str:
        """docx 渲染返回空串（必须走 render_to_file 落盘）"""
        return ""

    def render_to_file(self, data: RenderData, file_path: str) -> str:
        abs_path = os.path.abspath(file_path)
        os.makedirs(os.path.dirname(abs_path), exist_ok=True)

        doc = Document()
        judge = data.judge or {}
        sections = judge.get("sections") or {}
        metric = data.metric or {}
        alert = metric.get("alert") or {}

        # 标题
        title = data.extra.get("title", f"{data.cycle_label}网络安全态势报告")
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        meta = doc.add_paragraph()
        meta.add_run(
            f"周期：{data.cycle_label} ｜ 统计窗口：{data.window_start} ～ {data.window_end}\n"
            f"生成时间：{data.generated_at} ｜ 综合风险等级：{judge.get('risk_level', 'LOW')}"
        ).font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x90, 0x90, 0x90)

        # 章节正文
        section_titles = {
            "overview": "一、总体态势",
            "alert": "二、告警分析",
            "vuln": "三、漏洞与资产风险",
            "attack": "四、攻击行为研判",
            "trend": "五、趋势与预测",
            "suggestion": "六、安全建议",
        }
        for key, stitle in section_titles.items():
            doc.add_heading(stitle, level=1)
            text = (sections.get(key) or "").strip()
            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith(("1.", "2.", "3.", "4.", "5.", "6.")):
                    doc.add_paragraph(line, style="List Number")
                else:
                    doc.add_paragraph(line)

        # 告警等级表
        doc.add_heading("附：告警等级分布", level=2)
        table = doc.add_table(rows=2, cols=4)
        table.style = "Light Grid Accent 1"
        headers = ["高危", "中危", "低危", "提示"]
        values = [alert.get("high", 0), alert.get("medium", 0),
                  alert.get("low", 0), alert.get("info", 0)]
        for i, (hd, val) in enumerate(zip(headers, values)):
            table.rows[0].cells[i].text = hd
            table.rows[1].cells[i].text = str(val)

        doc.add_paragraph()
        doc.add_paragraph("本报告由多周期网安报告智能体自动生成，数据来源为安全运营平台采集日志与资产台账。")

        doc.save(abs_path)
        logger.info(f"[RENDER] DOCX 已保存: {abs_path}")
        return abs_path
