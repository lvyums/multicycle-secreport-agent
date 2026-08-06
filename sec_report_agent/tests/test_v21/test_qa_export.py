"""V2.1 测试 — 报告智能问答 + 导出（md/docx）
覆盖：QA 接口 LLM 模式/降级模式/参数校验/404；导出 md/docx 200、非法 format、404。
"""

import io

import pytest


# ── 工具 ──

class FakeLLM:
    """LLM 假客户端：chat 成功返回固定答案"""

    async def chat(self, messages, temperature=None, timeout=None):
        return {"content": "根据报告统计，高危告警共 291 起，建议优先处置。", "success": True, "error": None}


class FailingLLM:
    """LLM 假客户端：chat 失败（触发降级提取）"""

    async def chat(self, messages, temperature=None, timeout=None):
        return {"content": None, "success": False, "error": "mock fail"}


def _qa_body(version_id: int = 1, question: str = "高危告警有多少起？") -> dict:
    return {"versionId": version_id, "question": question}


@pytest.fixture()
def seed_version(db_session):
    """插入一个带正文的报告版本，返回 version id"""
    from model.entity.entities import ReportVersion

    ver = ReportVersion(
        task_id=999001,
        cycle="MONTHLY",
        window_start="2025-10-01 00:00:00",
        window_end="2025-11-01 00:00:00",
        version_no=1,
        version_type="AI_DRAFT",
        status="DRAFT",
        title="月报网络安全态势报告（2025-10-01 至 2025-11-01）",
        content_md=(
            "# 月报网络安全态势报告（2025-10-01 至 2025-11-01）\n\n"
            "## 一、总体态势\n\n本周期共监测到安全事件 651 起，其中高危 291 起，事件闭环率 63.1%。\n\n"
            "## 二、告警分析\n\n告警总量 651 起，高危 291 起。\n"
            "- 事件类型分布：brute_force=190, web_attack=180\n"
            "- 闭环率：63.1%\n\n"
            "| 等级 | 数量 |\n| --- | --- |\n| 高危 | 291 |\n| 中危 | 277 |\n\n"
            "## 六、安全建议\n\n1. 优先处置高危告警；\n2. 对 TOP 攻击源实施封禁。\n"
        ),
        operator="system",
    )
    db_session.add(ver)
    db_session.commit()
    db_session.refresh(ver)
    return ver.id


@pytest.fixture()
def db_session(client):
    """从 client 依赖拿到的 Session（conftest 的 rbac 覆盖同源）"""
    from infra.db.session import SessionLocal

    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ── 智能问答 ──

def test_qa_llm_mode(client, seed_version, monkeypatch):
    """LLM 可用 → 返回 llm 模式 + LLM 答案"""
    from capability.judge.llm_factory import LLMFactory

    async def fake_get():
        return FakeLLM()

    monkeypatch.setattr(LLMFactory, "get_light_llm", staticmethod(fake_get))
    r = client.post("/api/report/qa", json=_qa_body(seed_version))
    assert r.status_code == 200
    data = r.json()["data"]
    assert data["mode"] == "llm"
    assert "291" in data["answer"]


def test_qa_fallback_mode(client, seed_version, monkeypatch):
    """LLM 失败 → 降级为正文提取（fallback 模式）"""
    from capability.judge.llm_factory import LLMFactory

    async def fake_get():
        return FailingLLM()

    monkeypatch.setattr(LLMFactory, "get_light_llm", staticmethod(fake_get))
    r = client.post("/api/report/qa", json=_qa_body(seed_version))
    assert r.status_code == 200
    data = r.json()["data"]
    assert data["mode"] == "fallback"
    assert "高危" in data["answer"] or "告警" in data["answer"]


def test_qa_fallback_prefers_overview(client, seed_version, monkeypatch):
    """降级：问数量/告警类问题优先命中总体态势或告警分析章节"""
    from capability.judge.llm_factory import LLMFactory

    async def fake_get():
        return FailingLLM()

    monkeypatch.setattr(LLMFactory, "get_light_llm", staticmethod(fake_get))
    r = client.post("/api/report/qa", json=_qa_body(seed_version, "高危告警有多少起？"))
    data = r.json()["data"]
    assert data["mode"] == "fallback"
    assert "总体态势" in data["answer"] or "告警分析" in data["answer"]
    assert "651" in data["answer"] or "291" in data["answer"]


def test_qa_param_missing(client, seed_version):
    """参数校验：缺 versionId / 缺 question → 1001"""
    r1 = client.post("/api/report/qa", json={"question": "什么情况"})
    assert r1.json()["code"] == 400
    r2 = client.post("/api/report/qa", json={"versionId": seed_version, "question": "  "})
    assert r2.json()["code"] == 400


def test_qa_version_not_found(client):
    """版本不存在 → 404"""
    r = client.post("/api/report/qa", json=_qa_body(999999))
    assert r.json()["code"] == 404


def test_qa_empty_content(client, db_session):
    """版本正文为空 → empty 模式提示"""
    from model.entity.entities import ReportVersion

    ver = ReportVersion(
        task_id=999002, cycle="DAILY", window_start="2025-01-01 00:00:00",
        window_end="2025-01-02 00:00:00", version_no=1, status="DRAFT",
        title="空报告", content_md="", operator="system",
    )
    db_session.add(ver)
    db_session.commit()
    db_session.refresh(ver)
    r = client.post("/api/report/qa", json=_qa_body(ver.id))
    data = r.json()["data"]
    assert data["mode"] == "empty"
    assert "无报告正文" in data["answer"]


# ── 导出 ──

def test_export_markdown(client, seed_version):
    """导出 md → 200 + 附件头 + 正文内容"""
    r = client.get(f"/api/report/export/{seed_version}?format=md")
    assert r.status_code == 200
    assert "text/markdown" in r.headers.get("content-type", "")
    assert "attachment" in r.headers.get("content-disposition", "")
    body = r.content.decode("utf-8")
    assert "一、总体态势" in body
    assert "651" in body


def test_export_docx(client, seed_version):
    """导出 docx → 200 + 可被 python-docx 解析 + 含标题"""
    from docx import Document

    r = client.get(f"/api/report/export/{seed_version}?format=docx")
    assert r.status_code == 200
    assert "attachment" in r.headers.get("content-disposition", "")
    doc = Document(io.BytesIO(r.content))
    texts = [p.text for p in doc.paragraphs]
    assert any("总体态势" in t for t in texts)
    assert any("高危" in t for t in texts)


def test_export_bad_format(client, seed_version):
    """非法 format → 422（FastAPI pattern 校验）"""
    r = client.get(f"/api/report/export/{seed_version}?format=pdf")
    assert r.status_code == 422


def test_export_not_found(client):
    """版本不存在 → 404"""
    r = client.get("/api/report/export/999999?format=md")
    assert r.json()["code"] == 404


def test_export_filename_chinese(client, seed_version):
    """中文标题文件名 → URL 编码后仍 200（防止 latin-1 500）"""
    r = client.get(f"/api/report/export/{seed_version}?format=md")
    assert r.status_code == 200
    cd = r.headers.get("content-disposition", "")
    assert "UTF-8''" in cd


# ── 导出服务单测（docx 表格/列表） ──

def test_build_docx_table_and_list():
    """docx 构建：markdown 表格 + 列表转 docx 结构"""
    from docx import Document
    from app.services.report_export_service import ReportExportService

    md = "# 标题\n\n## 章节\n\n- 列表项A\n- 列表项B\n\n| 等级 | 数量 |\n| --- | --- |\n| 高危 | 291 |\n"
    payload = ReportExportService.build_docx(md)
    doc = Document(io.BytesIO(payload))
    assert doc.paragraphs[0].text == "标题"
    assert doc.tables and len(doc.tables[0].rows) == 2
    assert doc.tables[0].rows[1].cells[0].text == "高危"
